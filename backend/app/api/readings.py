import csv
import io
import json
import statistics
from collections import Counter, defaultdict
from datetime import datetime, timedelta, timezone
from typing import Annotated, Any, Generator, Optional

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import JSONResponse, StreamingResponse
from sqlalchemy import func, text
from sqlalchemy.orm import Session

from ..config import get_settings
from ..db import get_db
from ..models import (
    Alert,
    BehaviourLabel,
    Cage,
    CageSocialSample,
    MouseStateSample,
    Reading,
    User,
)
from ..schemas import ReadingOut, SummaryDay, SummaryOut
from ..security import get_current_user

router = APIRouter(prefix="/api/v1/cages", tags=["readings"])

# Maximum rows for streaming exports (guards against unbounded 30-second queries)
_EXPORT_ROW_CAP = 50_000


def _ensure_cage(db: Session, cage_id: str) -> Cage:
    cage = db.query(Cage).filter(Cage.id == cage_id).one_or_none()
    if not cage:
        raise HTTPException(status_code=404, detail={"code": "NOT_FOUND", "message": "Unknown cage."})
    return cage


@router.get("/{cage_id}/readings", response_model=list[ReadingOut])
def list_readings(
    cage_id: str,
    db: Annotated[Session, Depends(get_db)],
    _user: Annotated[User, Depends(get_current_user)],
    sensor: Optional[str] = None,
    from_: Optional[datetime] = Query(default=None, alias="from"),
    to: Optional[datetime] = None,
    limit: int = Query(default=500, ge=1, le=5000),
) -> list[Reading]:
    _ensure_cage(db, cage_id)
    q = db.query(Reading).filter(Reading.cage_id == cage_id)
    if sensor:
        q = q.filter(Reading.sensor == sensor)
    if from_:
        q = q.filter(Reading.ts >= from_)
    if to:
        q = q.filter(Reading.ts <= to)
    return q.order_by(Reading.ts.desc()).limit(limit).all()


@router.get("/{cage_id}/summary", response_model=SummaryOut)
def summary(
    cage_id: str,
    db: Annotated[Session, Depends(get_db)],
    _user: Annotated[User, Depends(get_current_user)],
    days: int = Query(default=7, ge=1, le=90),
) -> SummaryOut:
    _ensure_cage(db, cage_id)
    end = datetime.now(tz=timezone.utc)
    start = end - timedelta(days=days)

    # ── Feeding: GROUP BY day — sum delivered_g, wasted_g; min remaining_g ──────
    # func.date() works on both SQLite (date string) and Postgres (date type).
    # JSON field extraction uses json_extract (SQLite) / ->> (Postgres); we use
    # SQLAlchemy text() for the JSON path so it stays portable within our target
    # DBs. The payload column stores JSON as text in SQLite and as JSONB in Postgres,
    # but since we only target SQLite in prod, json_extract is correct here.
    # For cross-DB portability we choose a Python-side aggregation approach but
    # push date bucketing + per-sensor filtering to the DB layer via SQL.

    # Feeding rows — only pull the columns we need, grouped by day in Python
    # after fetching per-sensor with indexed WHERE sensor=X clause.
    feeding_rows = (
        db.query(
            func.date(Reading.ts).label("day"),
            func.sum(text("CAST(json_extract(readings.payload, '$.delivered_g') AS REAL)")).label("delivered_sum"),
            func.sum(text("CAST(json_extract(readings.payload, '$.wasted_g') AS REAL)")).label("wasted_sum"),
            func.min(text("CAST(json_extract(readings.payload, '$.remaining_g') AS REAL)")).label("remaining_min"),
        )
        .filter(
            Reading.cage_id == cage_id,
            Reading.sensor == "feeding",
            Reading.ts >= start,
            Reading.ts <= end,
        )
        .group_by(func.date(Reading.ts))
        .all()
    )

    # Water rows — sum flow_ml per day
    water_rows = (
        db.query(
            func.date(Reading.ts).label("day"),
            func.sum(text("CAST(json_extract(readings.payload, '$.flow_ml') AS REAL)")).label("flow_sum"),
        )
        .filter(
            Reading.cage_id == cage_id,
            Reading.sensor == "water",
            Reading.ts >= start,
            Reading.ts <= end,
        )
        .group_by(func.date(Reading.ts))
        .all()
    )

    # Metabolic rows — avg rer, sum ee_kcal_h (converted to kcal), count per day
    metabolic_rows = (
        db.query(
            func.date(Reading.ts).label("day"),
            func.avg(text("CAST(json_extract(readings.payload, '$.rer') AS REAL)")).label("rer_avg"),
            func.sum(text("CAST(json_extract(readings.payload, '$.ee_kcal_h') AS REAL)")).label("ee_sum"),
            func.count(Reading.id).label("n"),
        )
        .filter(
            Reading.cage_id == cage_id,
            Reading.sensor == "metabolic",
            Reading.ts >= start,
            Reading.ts <= end,
        )
        .group_by(func.date(Reading.ts))
        .all()
    )

    # Behaviour labels — count per day per label
    label_rows = (
        db.query(
            func.date(BehaviourLabel.ts).label("day"),
            BehaviourLabel.label,
            func.count(BehaviourLabel.id).label("cnt"),
        )
        .filter(
            BehaviourLabel.cage_id == cage_id,
            BehaviourLabel.ts >= start,
        )
        .group_by(func.date(BehaviourLabel.ts), BehaviourLabel.label)
        .all()
    )

    # Collect all unique days across all sensor types
    all_days: set[str] = set()
    for r in feeding_rows:
        all_days.add(str(r.day))
    for r in water_rows:
        all_days.add(str(r.day))
    for r in metabolic_rows:
        all_days.add(str(r.day))
    for r in label_rows:
        all_days.add(str(r.day))

    # Build lookup dicts keyed by ISO day string
    feeding_by_day: dict[str, dict] = {}
    for r in feeding_rows:
        day = str(r.day)
        feeding_by_day[day] = {
            "delivered_g": float(r.delivered_sum or 0.0),
            "wasted_g": float(r.wasted_sum or 0.0),
            "remaining_min_g": float(r.remaining_min) if r.remaining_min is not None else 0.0,
        }

    water_by_day: dict[str, float] = {}
    for r in water_rows:
        water_by_day[str(r.day)] = float(r.flow_sum or 0.0)

    metabolic_by_day: dict[str, dict] = {}
    for r in metabolic_rows:
        day = str(r.day)
        # ee_kcal_h * (5 / 3600) matches original: 5-second measurement interval
        ee_sum_raw = float(r.ee_sum or 0.0)
        n = int(r.n or 0)
        metabolic_by_day[day] = {
            "rer_avg": float(r.rer_avg or 0.0),
            "ee_kcal": ee_sum_raw * (5.0 / 3600.0),
        }

    label_by_day: dict[str, dict[str, int]] = {}
    for r in label_rows:
        day = str(r.day)
        d = label_by_day.setdefault(day, {})
        d[r.label] = int(r.cnt)

    # Assemble response — preserve exact shape the frontend expects
    out: list[SummaryDay] = []
    for day in sorted(all_days):
        feeding = feeding_by_day.get(day, {"delivered_g": 0.0, "wasted_g": 0.0, "remaining_min_g": 0.0})
        delivered = feeding["delivered_g"]
        wasted = feeding["wasted_g"]
        intake = max(delivered - wasted, 0.0)

        metabolic = metabolic_by_day.get(day, {"rer_avg": 0.0, "ee_kcal": 0.0})

        lab = label_by_day.get(day, {})
        out.append(SummaryDay(
            date=day,
            feeding={
                "intake_g": round(intake, 2),
                "delivered_g": round(delivered, 2),
                "wasted_g": round(wasted, 2),
            },
            water_ml=round(water_by_day.get(day, 0.0), 2),
            metabolic={
                "avg_rer": round(metabolic["rer_avg"], 3),
                "ee_kcal": round(metabolic["ee_kcal"], 2),
            },
            behaviour={
                "sleeping_minutes": lab.get("sleeping", 0),
                "active_minutes": sum(v for k, v in lab.items() if k != "sleeping"),
            },
        ))

    return SummaryOut(cage_id=cage_id, days=out)


# ─── Flexible time-series endpoint (interactive Reports) ─────────────────────

# Available metrics: maps a friendly key → (sensor, payload field, unit, label).
# The frontend's metric dropdown is built from /timeseries/metrics.
_METRIC_CATALOG: dict[str, dict] = {
    "temperature_c":  {"sensor": "environment", "field": "temperature_c", "unit": "°C",   "label": "Temperature",   "default_agg": "avg"},
    "humidity_pct":   {"sensor": "environment", "field": "humidity_pct",  "unit": "%",    "label": "Humidity",      "default_agg": "avg"},
    "co2_ppm":        {"sensor": "environment", "field": "co2_ppm",       "unit": "ppm",  "label": "CO₂",            "default_agg": "avg"},
    "o2_pct":         {"sensor": "environment", "field": "o2_pct",        "unit": "%",    "label": "O₂",             "default_agg": "avg"},
    "animal_g":       {"sensor": "weighing",    "field": "animal_g",      "unit": "g",    "label": "Animal weight",  "default_agg": "avg"},
    "delivered_g":    {"sensor": "feeding",     "field": "delivered_g",   "unit": "g",    "label": "Food delivered", "default_agg": "sum"},
    "remaining_g":    {"sensor": "feeding",     "field": "remaining_g",   "unit": "g",    "label": "Food remaining", "default_agg": "last"},
    "wasted_g":       {"sensor": "feeding",     "field": "wasted_g",      "unit": "g",    "label": "Food wasted",    "default_agg": "sum"},
    "flow_ml":        {"sensor": "water",       "field": "flow_ml",       "unit": "mL",   "label": "Water flow",     "default_agg": "sum"},
    "tank_ml":        {"sensor": "water",       "field": "tank_ml",       "unit": "mL",   "label": "Water tank",     "default_agg": "last"},
    "rer":            {"sensor": "metabolic",   "field": "rer",           "unit": "",     "label": "RER",            "default_agg": "avg"},
    "vo2_ml_min_kg":  {"sensor": "metabolic",   "field": "vo2_ml_min_kg", "unit": "mL/min/kg", "label": "VO₂",       "default_agg": "avg"},
    "vco2_ml_min_kg": {"sensor": "metabolic",   "field": "vco2_ml_min_kg","unit": "mL/min/kg", "label": "VCO₂",      "default_agg": "avg"},
    "ee_kcal_h":      {"sensor": "metabolic",   "field": "ee_kcal_h",     "unit": "kcal/h","label": "Energy expend.","default_agg": "avg"},
    "movement_cm":    {"sensor": "behaviour",   "field": "movement_cm",   "unit": "cm",   "label": "Movement",       "default_agg": "sum"},
}

_GRANULARITY_SECONDS = {
    "second": 1,
    "minute": 60,
    "hour":   3600,
    "day":    86400,
}


@router.get("/timeseries/metrics")
def timeseries_metrics(
    _user: Annotated[User, Depends(get_current_user)],
) -> dict:
    """List every metric the timeseries endpoint can chart (for the UI dropdown)."""
    return {
        "metrics": [
            {"key": k, "label": v["label"], "unit": v["unit"],
             "sensor": v["sensor"], "default_agg": v["default_agg"]}
            for k, v in _METRIC_CATALOG.items()
        ],
        "granularities": list(_GRANULARITY_SECONDS.keys()),
        "aggregations": ["avg", "sum", "min", "max", "last", "count"],
    }


def _bucket_key(ts: datetime, granularity: str) -> datetime:
    """Truncate a timestamp down to the start of its granularity bucket."""
    if granularity == "second":
        return ts.replace(microsecond=0)
    if granularity == "minute":
        return ts.replace(second=0, microsecond=0)
    if granularity == "hour":
        return ts.replace(minute=0, second=0, microsecond=0)
    return ts.replace(hour=0, minute=0, second=0, microsecond=0)  # day


@router.get("/{cage_id}/timeseries")
def timeseries(
    cage_id: str,
    db: Annotated[Session, Depends(get_db)],
    _user: Annotated[User, Depends(get_current_user)],
    metric: str = Query(..., description="metric key from /timeseries/metrics"),
    granularity: str = Query(default="hour"),
    agg: str = Query(default="avg"),
    window: int = Query(default=24, ge=1, le=10000,
                        description="number of granularity buckets to look back"),
) -> dict:
    """Bucketed, aggregated time-series for ONE metric — powers the interactive
    Reports charts. Pick metric + granularity (second/minute/hour/day) +
    aggregation (avg/sum/min/max/last/count) + how many buckets to look back."""
    _ensure_cage(db, cage_id)

    spec = _METRIC_CATALOG.get(metric)
    if spec is None:
        raise HTTPException(status_code=400,
                            detail={"code": "BAD_METRIC", "message": f"Unknown metric '{metric}'."})
    if granularity not in _GRANULARITY_SECONDS:
        raise HTTPException(status_code=400,
                            detail={"code": "BAD_GRANULARITY", "message": "granularity must be second/minute/hour/day."})
    if agg not in ("avg", "sum", "min", "max", "last", "count"):
        raise HTTPException(status_code=400,
                            detail={"code": "BAD_AGG", "message": "agg must be avg/sum/min/max/last/count."})

    span_s = _GRANULARITY_SECONDS[granularity] * window
    end = datetime.now(tz=timezone.utc)
    start = end - timedelta(seconds=span_s)

    rows = (
        db.query(Reading)
        .filter(Reading.cage_id == cage_id,
                Reading.sensor == spec["sensor"],
                Reading.ts >= start)
        .order_by(Reading.ts)
        .all()
    )

    field = spec["field"]
    buckets: dict[datetime, list[float]] = {}
    for r in rows:
        if not r.payload or field not in r.payload:
            continue
        try:
            val = float(r.payload[field])
        except (TypeError, ValueError):
            continue
        buckets.setdefault(_bucket_key(r.ts, granularity), []).append(val)

    def _aggregate(vals: list[float]) -> float:
        if agg == "avg":   return sum(vals) / len(vals)
        if agg == "sum":   return sum(vals)
        if agg == "min":   return min(vals)
        if agg == "max":   return max(vals)
        if agg == "last":  return vals[-1]
        return float(len(vals))  # count

    points = [
        {"t": bk.isoformat().replace("+00:00", "Z"),
         "v": round(_aggregate(vals), 4),
         "n": len(vals)}
        for bk, vals in sorted(buckets.items())
    ]

    # Quick stats for the UI's KPI row
    all_vals = [p["v"] for p in points]
    stats = {
        "count": len(points),
        "min": round(min(all_vals), 4) if all_vals else None,
        "max": round(max(all_vals), 4) if all_vals else None,
        "avg": round(sum(all_vals) / len(all_vals), 4) if all_vals else None,
        "first": all_vals[0] if all_vals else None,
        "last": all_vals[-1] if all_vals else None,
    }

    return {
        "cage_id": cage_id,
        "metric": metric,
        "label": spec["label"],
        "unit": spec["unit"],
        "granularity": granularity,
        "agg": agg,
        "window": window,
        "from": start.isoformat().replace("+00:00", "Z"),
        "to": end.isoformat().replace("+00:00", "Z"),
        "stats": stats,
        "points": points,
    }


# ─── CSV export — streaming, date-windowed ───────────────────────────────────

@router.get("/{cage_id}/export.csv")
def export_csv(
    cage_id: str,
    db: Annotated[Session, Depends(get_db)],
    _user: Annotated[User, Depends(get_current_user)],
    from_: Optional[datetime] = Query(default=None, alias="from"),
    to: Optional[datetime] = None,
    days: Optional[int] = Query(default=None, ge=1, le=365,
                                description="trailing-day window (ignored when 'from' is supplied)"),
    limit: int = Query(default=_EXPORT_ROW_CAP, ge=1, le=_EXPORT_ROW_CAP,
                       description="max rows returned"),
):
    """Stream CSV rows incrementally. Honors from/to params as before; when
    neither is provided applies a default 30-day trailing window so the export
    can never run unbounded."""
    _ensure_cage(db, cage_id)

    # Resolve date window — default to 30 days if no explicit range is given
    if from_ is None:
        window_days = days if days is not None else 30
        from_ = datetime.now(tz=timezone.utc) - timedelta(days=window_days)

    q = db.query(Reading).filter(Reading.cage_id == cage_id)
    q = q.filter(Reading.ts >= from_)
    if to:
        q = q.filter(Reading.ts <= to)
    q = q.order_by(Reading.ts).limit(limit)

    # Two-pass approach: first fetch all payload keys (needed for CSV header),
    # then stream the rows. We use yield_per so rows are not all held in memory.
    # Collecting only keys requires a light scan; we cap it the same way.
    key_scan = q.with_entities(Reading.payload).yield_per(1000)
    all_payload_keys: list[str] = []
    for (payload,) in key_scan:
        if payload:
            for k in payload:
                if k not in all_payload_keys:
                    all_payload_keys.append(k)

    def _generate_csv() -> Generator[str, None, None]:
        buf = io.StringIO()
        writer = csv.writer(buf)
        writer.writerow(["ts", "cage_id", "sensor", "seq"] + all_payload_keys)
        yield buf.getvalue()
        buf.seek(0)
        buf.truncate(0)

        row_q = db.query(Reading).filter(Reading.cage_id == cage_id)
        row_q = row_q.filter(Reading.ts >= from_)
        if to:
            row_q = row_q.filter(Reading.ts <= to)
        row_q = row_q.order_by(Reading.ts).limit(limit)

        for r in row_q.yield_per(500):
            flat = [r.payload.get(k, "") if r.payload else "" for k in all_payload_keys]
            writer.writerow([r.ts.isoformat(), r.cage_id, r.sensor, r.seq] + flat)
            if buf.tell() >= 65536:  # flush every ~64 KB
                yield buf.getvalue()
                buf.seek(0)
                buf.truncate(0)

        remainder = buf.getvalue()
        if remainder:
            yield remainder

    return StreamingResponse(
        _generate_csv(),
        media_type="text/csv",
        headers={"Content-Disposition": f"attachment; filename={cage_id}.csv"},
    )


# ─── Multi-format exports ────────────────────────────────────────────────────


def _collect_export_data(db: Session, cage: Cage, days: int) -> dict:
    """Pull readings + labels + alerts for the trailing `days` days.

    Row count is capped at _EXPORT_ROW_CAP so large exports don't OOM.
    """
    end = datetime.now(tz=timezone.utc)
    start = end - timedelta(days=days)
    readings = (
        db.query(Reading)
        .filter(Reading.cage_id == cage.id, Reading.ts >= start)
        .order_by(Reading.ts)
        .limit(_EXPORT_ROW_CAP)
        .all()
    )
    labels = (
        db.query(BehaviourLabel)
        .filter(BehaviourLabel.cage_id == cage.id, BehaviourLabel.ts >= start)
        .order_by(BehaviourLabel.ts)
        .all()
    )
    alerts = (
        db.query(Alert)
        .filter(Alert.cage_id == cage.id, Alert.ts >= start)
        .order_by(Alert.ts)
        .all()
    )
    return {
        "cage": cage,
        "from": start,
        "to": end,
        "readings": readings,
        "labels": labels,
        "alerts": alerts,
    }


@router.get("/{cage_id}/export.xlsx")
def export_xlsx(
    cage_id: str,
    db: Annotated[Session, Depends(get_db)],
    _user: Annotated[User, Depends(get_current_user)],
    days: int = Query(default=7, ge=1, le=90),
):
    """Excel workbook with formatted sheets: Summary, Readings, Labels, Alerts."""
    cage = _ensure_cage(db, cage_id)
    data = _collect_export_data(db, cage, days)

    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font, PatternFill
        from openpyxl.utils import get_column_letter
    except ImportError:
        raise HTTPException(
            status_code=503,
            detail={"code": "MISSING_DEP", "message": "openpyxl not installed on server."},
        )

    wb = Workbook()
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill(start_color="0F172A", end_color="0F172A", fill_type="solid")
    title_font = Font(bold=True, size=14, color="10B981")

    # 1. Summary sheet
    s = wb.active
    s.title = "Summary"
    s["A1"] = f"IVC Cage Export — {cage.name}"
    s["A1"].font = title_font
    s["A2"] = f"Cage ID: {cage.id}"
    s["A3"] = f"Strain: {cage.animal_strain or '—'}"
    s["A4"] = f"Range: {data['from'].date()}  →  {data['to'].date()}  ({days} days)"
    s["A5"] = f"Generated: {datetime.now(tz=timezone.utc).isoformat(timespec='seconds')}"
    s["A7"] = "Counts"
    s["A7"].font = Font(bold=True, size=12)
    s["A8"], s["B8"] = "Readings", len(data["readings"])
    s["A9"], s["B9"] = "Behaviour labels", len(data["labels"])
    s["A10"], s["B10"] = "Alerts", len(data["alerts"])
    for col in ("A", "B"):
        s.column_dimensions[col].width = 28

    # 2. Readings sheet — flatten payload fields into separate columns
    ws = wb.create_sheet("Readings")
    # Collect all unique payload keys across all readings
    all_payload_keys: list[str] = []
    for r in data["readings"]:
        if r.payload:
            for k in r.payload:
                if k not in all_payload_keys:
                    all_payload_keys.append(k)
    # Human-readable header names: delivered_g → Delivered (g), temperature_c → Temperature (°C), etc.
    _FRIENDLY: dict[str, str] = {
        "delivered_g": "Delivered (g)", "remaining_g": "Remaining (g)", "wasted_g": "Wasted (g)",
        "gate_state": "Gate state", "flow_ml": "Flow (mL)", "tank_ml": "Tank (mL)",
        "vo2_ml_min_kg": "VO₂ (mL/min/kg)", "vco2_ml_min_kg": "VCO₂ (mL/min/kg)",
        "rer": "RER", "ee_kcal_h": "EE (kcal/h)",
        "grid_cells": "Grid cells", "movement_cm": "Movement (cm)",
        "animal_g": "Animal weight (g)",
        "temperature_c": "Temperature (°C)", "humidity_pct": "Humidity (%)",
        "co2_ppm": "CO₂ (ppm)", "o2_pct": "O₂ (%)",
    }
    friendly_keys = [_FRIENDLY.get(k, k) for k in all_payload_keys]
    headers = ["Timestamp (UTC)", "Sensor", "Seq"] + friendly_keys
    ws.append(headers)
    for ci, _ in enumerate(headers, 1):
        c = ws.cell(row=1, column=ci)
        c.font, c.fill, c.alignment = header_font, header_fill, Alignment(horizontal="center")
    for r in data["readings"]:
        flat = []
        for k in all_payload_keys:
            v = r.payload.get(k, "") if r.payload else ""
            # Convert lists (grid_cells) to string for display
            if isinstance(v, list):
                v = ", ".join(str(x) for x in v)
            flat.append(v)
        ws.append([r.ts.isoformat(timespec="seconds"), r.sensor, r.seq] + flat)
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = f"A1:{get_column_letter(len(headers))}{max(ws.max_row, 1)}"
    base_widths = [22, 14, 8]
    payload_widths = [max(14, len(fk) + 2) for fk in friendly_keys]
    for ci, width in enumerate(base_widths + payload_widths, 1):
        ws.column_dimensions[get_column_letter(ci)].width = width

    # 3. Behaviour labels sheet
    lws = wb.create_sheet("Behaviour")
    lheaders = ["Timestamp (UTC)", "Label", "Confidence"]
    lws.append(lheaders)
    for ci, _ in enumerate(lheaders, 1):
        c = lws.cell(row=1, column=ci)
        c.font, c.fill, c.alignment = header_font, header_fill, Alignment(horizontal="center")
    for lab in data["labels"]:
        lws.append([lab.ts.isoformat(timespec="seconds"), lab.label, round(lab.confidence, 4)])
    lws.freeze_panes = "A2"
    for ci, width in enumerate([22, 18, 14], 1):
        lws.column_dimensions[get_column_letter(ci)].width = width

    # 4. Alerts sheet
    aws = wb.create_sheet("Alerts")
    aheaders = ["Timestamp (UTC)", "Severity", "Rule", "Message", "Acknowledged"]
    aws.append(aheaders)
    for ci, _ in enumerate(aheaders, 1):
        c = aws.cell(row=1, column=ci)
        c.font, c.fill, c.alignment = header_font, header_fill, Alignment(horizontal="center")
    for a in data["alerts"]:
        aws.append([
            a.ts.isoformat(timespec="seconds"),
            a.severity,
            a.rule,
            a.message,
            a.acknowledged_at.isoformat(timespec="seconds") if a.acknowledged_at else "",
        ])
    aws.freeze_panes = "A2"
    for ci, width in enumerate([22, 12, 24, 60, 22], 1):
        aws.column_dimensions[get_column_letter(ci)].width = width

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    fname = f"{cage_id}_{datetime.now().strftime('%Y%m%d')}.xlsx"
    return StreamingResponse(
        iter([buf.getvalue()]),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={fname}"},
    )


@router.get("/{cage_id}/export.pdf")
def export_pdf(
    cage_id: str,
    db: Annotated[Session, Depends(get_db)],
    _user: Annotated[User, Depends(get_current_user)],
    days: int = Query(default=7, ge=1, le=90),
):
    """PDF report with summary KPIs + recent readings + alerts table."""
    cage = _ensure_cage(db, cage_id)
    data = _collect_export_data(db, cage, days)

    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import (
            SimpleDocTemplate,
            Paragraph,
            Spacer,
            Table,
            TableStyle,
            PageBreak,
        )
    except ImportError:
        raise HTTPException(
            status_code=503,
            detail={"code": "MISSING_DEP", "message": "reportlab not installed on server."},
        )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=1.8 * cm, rightMargin=1.8 * cm,
        topMargin=1.8 * cm, bottomMargin=1.8 * cm,
        title=f"IVC Cage Report — {cage.name}",
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("title", parent=styles["Title"], textColor=colors.HexColor("#10b981"))
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], textColor=colors.HexColor("#0f172a"))
    body = styles["BodyText"]

    story = []
    story.append(Paragraph(f"IVC Cage Report — {cage.name}", title_style))
    story.append(Paragraph(f"<b>Cage ID:</b> {cage.id}  &nbsp; <b>Strain:</b> {cage.animal_strain or '—'}", body))
    story.append(Paragraph(
        f"<b>Range:</b> {data['from'].date()} → {data['to'].date()} ({days} days)  &nbsp; "
        f"<b>Generated:</b> {datetime.now(tz=timezone.utc).isoformat(timespec='seconds')}",
        body,
    ))
    story.append(Spacer(1, 0.4 * cm))

    # KPI block
    feeding_total, water_total, rer_vals, weight_vals = 0.0, 0.0, [], []
    for r in data["readings"]:
        if r.sensor == "feeding":
            feeding_total += float(r.payload.get("delivered_g", 0))
        elif r.sensor == "water":
            water_total += float(r.payload.get("flow_ml", 0))
        elif r.sensor == "metabolic":
            v = r.payload.get("rer")
            if v is not None:
                rer_vals.append(float(v))
        elif r.sensor == "weighing":
            v = r.payload.get("animal_g")
            if v is not None:
                weight_vals.append(float(v))
    avg_rer = sum(rer_vals) / len(rer_vals) if rer_vals else 0.0
    last_weight = weight_vals[-1] if weight_vals else 0.0

    kpi_rows = [
        ["Total feed delivered", f"{feeding_total:.1f} g"],
        ["Total water flow", f"{water_total:.1f} mL"],
        ["Average RER", f"{avg_rer:.3f}"],
        ["Latest weight", f"{last_weight:.1f} g"],
        ["Readings count", str(len(data["readings"]))],
        ["Behaviour labels", str(len(data["labels"]))],
        ["Alerts", str(len(data["alerts"]))],
    ]
    kt = Table(kpi_rows, colWidths=[7 * cm, 5 * cm])
    kt.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f1f5f9")),
        ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#0f172a")),
        ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("ROWBACKGROUNDS", (0, 0), (-1, -1), [colors.HexColor("#f8fafc"), colors.HexColor("#f1f5f9")]),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd5e1")),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(Paragraph("Summary KPIs", h2))
    story.append(kt)
    story.append(Spacer(1, 0.5 * cm))

    # Alerts table
    if data["alerts"]:
        story.append(Paragraph("Alerts", h2))
        a_rows = [["Time (UTC)", "Severity", "Rule", "Message"]]
        for a in data["alerts"][-20:]:
            a_rows.append([
                a.ts.strftime("%Y-%m-%d %H:%M"),
                a.severity,
                a.rule,
                (a.message[:60] + "…") if len(a.message) > 60 else a.message,
            ])
        at = Table(a_rows, colWidths=[3.4 * cm, 1.8 * cm, 3 * cm, 8 * cm])
        at.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0f172a")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8.5),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
            ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#cbd5e1")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ]))
        story.append(at)
        story.append(Spacer(1, 0.4 * cm))

    # Recent readings (last 50)
    story.append(PageBreak())
    story.append(Paragraph("Recent readings (last 50)", h2))
    if data["readings"]:
        r_rows = [["Time (UTC)", "Sensor", "Seq", "Payload"]]
        for r in data["readings"][-50:]:
            payload_str = json.dumps(r.payload, separators=(",", ":"))
            r_rows.append([
                r.ts.strftime("%Y-%m-%d %H:%M"),
                r.sensor,
                str(r.seq),
                (payload_str[:80] + "…") if len(payload_str) > 80 else payload_str,
            ])
        rt = Table(r_rows, colWidths=[3.4 * cm, 2.4 * cm, 1.4 * cm, 9 * cm])
        rt.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0f172a")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 7.5),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
            ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#cbd5e1")),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(rt)

    doc.build(story)
    buf.seek(0)
    fname = f"{cage_id}_{datetime.now().strftime('%Y%m%d')}.pdf"
    return StreamingResponse(
        iter([buf.getvalue()]),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={fname}"},
    )


@router.get("/{cage_id}/export.docx")
def export_docx(
    cage_id: str,
    db: Annotated[Session, Depends(get_db)],
    _user: Annotated[User, Depends(get_current_user)],
    days: int = Query(default=7, ge=1, le=90),
):
    """Word document — readable report for non-technical stakeholders."""
    cage = _ensure_cage(db, cage_id)
    data = _collect_export_data(db, cage, days)

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(
            status_code=503,
            detail={"code": "MISSING_DEP", "message": "python-docx not installed on server."},
        )

    doc = Document()

    title = doc.add_heading(f"IVC Cage Report — {cage.name}", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)

    p = doc.add_paragraph()
    p.add_run("Cage ID: ").bold = True
    p.add_run(f"{cage.id}    ")
    p.add_run("Strain: ").bold = True
    p.add_run(f"{cage.animal_strain or '—'}")

    p = doc.add_paragraph()
    p.add_run("Range: ").bold = True
    p.add_run(f"{data['from'].date()} → {data['to'].date()} ({days} days)    ")
    p.add_run("Generated: ").bold = True
    p.add_run(f"{datetime.now(tz=timezone.utc).isoformat(timespec='seconds')}")

    # KPIs
    feeding_total, water_total, rer_vals, weight_vals = 0.0, 0.0, [], []
    for r in data["readings"]:
        if r.sensor == "feeding":
            feeding_total += float(r.payload.get("delivered_g", 0))
        elif r.sensor == "water":
            water_total += float(r.payload.get("flow_ml", 0))
        elif r.sensor == "metabolic":
            v = r.payload.get("rer")
            if v is not None:
                rer_vals.append(float(v))
        elif r.sensor == "weighing":
            v = r.payload.get("animal_g")
            if v is not None:
                weight_vals.append(float(v))
    avg_rer = sum(rer_vals) / len(rer_vals) if rer_vals else 0.0
    last_weight = weight_vals[-1] if weight_vals else 0.0

    doc.add_heading("Summary KPIs", level=1)
    kt = doc.add_table(rows=0, cols=2)
    kt.style = "Light Grid Accent 1"
    for label, val in [
        ("Total feed delivered", f"{feeding_total:.1f} g"),
        ("Total water flow", f"{water_total:.1f} mL"),
        ("Average RER", f"{avg_rer:.3f}"),
        ("Latest weight", f"{last_weight:.1f} g"),
        ("Readings count", str(len(data["readings"]))),
        ("Behaviour labels", str(len(data["labels"]))),
        ("Alerts", str(len(data["alerts"]))),
    ]:
        row = kt.add_row()
        row.cells[0].text = label
        row.cells[1].text = val
        row.cells[0].paragraphs[0].runs[0].bold = True

    # Alerts
    if data["alerts"]:
        doc.add_heading("Alerts", level=1)
        at = doc.add_table(rows=1, cols=4)
        at.style = "Medium Grid 1 Accent 1"
        hdr = at.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Time (UTC)", "Severity", "Rule", "Message"
        for a in data["alerts"][-30:]:
            row = at.add_row().cells
            row[0].text = a.ts.strftime("%Y-%m-%d %H:%M")
            row[1].text = a.severity
            row[2].text = a.rule
            row[3].text = a.message

    # Recent readings
    doc.add_heading("Recent readings (last 50)", level=1)
    if data["readings"]:
        rt = doc.add_table(rows=1, cols=4)
        rt.style = "Light List Accent 1"
        hdr = rt.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Time (UTC)", "Sensor", "Seq", "Payload"
        for r in data["readings"][-50:]:
            row = rt.add_row().cells
            row[0].text = r.ts.strftime("%Y-%m-%d %H:%M")
            row[1].text = r.sensor
            row[2].text = str(r.seq)
            payload = json.dumps(r.payload, separators=(",", ":"))
            row[3].text = (payload[:120] + "…") if len(payload) > 120 else payload

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    fname = f"{cage_id}_{datetime.now().strftime('%Y%m%d')}.docx"
    return StreamingResponse(
        iter([buf.getvalue()]),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={fname}"},
    )


@router.get("/{cage_id}/export.json")
def export_json(
    cage_id: str,
    db: Annotated[Session, Depends(get_db)],
    _user: Annotated[User, Depends(get_current_user)],
    days: int = Query(default=7, ge=1, le=90),
    from_: Optional[datetime] = Query(default=None, alias="from"),
    to: Optional[datetime] = None,
    limit: int = Query(default=_EXPORT_ROW_CAP, ge=1, le=_EXPORT_ROW_CAP,
                       description="max rows returned"),
):
    """Streaming JSON export — incremental chunked output for large datasets."""
    cage = _ensure_cage(db, cage_id)

    # Resolve date window
    end = to or datetime.now(tz=timezone.utc)
    start = from_ or (end - timedelta(days=days))

    labels = (
        db.query(BehaviourLabel)
        .filter(BehaviourLabel.cage_id == cage.id, BehaviourLabel.ts >= start)
        .order_by(BehaviourLabel.ts)
        .all()
    )
    alerts = (
        db.query(Alert)
        .filter(Alert.cage_id == cage.id, Alert.ts >= start)
        .order_by(Alert.ts)
        .all()
    )

    # Count total readings for the header (cheap COUNT query)
    reading_count = (
        db.query(func.count(Reading.id))
        .filter(Reading.cage_id == cage.id, Reading.ts >= start, Reading.ts <= end)
        .scalar()
    )
    capped_count = min(int(reading_count or 0), limit)

    def _generate_json() -> Generator[str, None, None]:
        header = {
            "schema_version": 1,
            "exported_at": datetime.now(tz=timezone.utc).isoformat(),
            "cage": {
                "id": cage.id,
                "name": cage.name,
                "animal_strain": cage.animal_strain,
                "animal_age_days": cage.animal_age_days,
                "location": cage.location,
            },
            "range": {
                "from": start.isoformat(),
                "to": end.isoformat(),
                "days": days,
            },
            "counts": {
                "readings": capped_count,
                "labels": len(labels),
                "alerts": len(alerts),
            },
        }
        yield json.dumps(header)[:-1]  # strip trailing "}" — we'll append arrays

        # Stream readings array
        yield ', "readings": ['
        first = True
        readings_q = (
            db.query(Reading)
            .filter(Reading.cage_id == cage.id, Reading.ts >= start, Reading.ts <= end)
            .order_by(Reading.ts)
            .limit(limit)
        )
        for r in readings_q.yield_per(500):
            row_json = json.dumps({"ts": r.ts.isoformat(), "sensor": r.sensor, "seq": r.seq, "payload": r.payload})
            if not first:
                yield ","
            yield row_json
            first = False
        yield "]"

        # Behaviour labels (already loaded — small dataset)
        yield ', "behaviour_labels": ' + json.dumps([
            {"ts": lab.ts.isoformat(), "label": lab.label, "confidence": lab.confidence}
            for lab in labels
        ])

        # Alerts (already loaded — small dataset)
        yield ', "alerts": ' + json.dumps([
            {
                "ts": a.ts.isoformat(),
                "severity": a.severity,
                "rule": a.rule,
                "message": a.message,
                "acknowledged_at": a.acknowledged_at.isoformat() if a.acknowledged_at else None,
            }
            for a in alerts
        ])

        yield "}"  # close the root object

    return StreamingResponse(
        _generate_json(),
        media_type="application/json",
        headers={"Content-Disposition": f"attachment; filename={cage_id}.json"},
    )


# ══════════════════════════════════════════════════════════════════════════════
#  Camera-driven behavioural reports — per-mouse time-in-state
#  Source: mouse_state_samples / cage_social_samples written by the camera
#  pipeline (BehavioralMonitor). seconds = COUNT(samples) * sample_interval_s.
# ══════════════════════════════════════════════════════════════════════════════

_CAM_STATES = ["Wake / Active", "Quiet / Rest", "Possible Sleep"]


def _sample_interval_s() -> float:
    return get_settings().camera_sample_interval_s


def _since(days: int) -> datetime:
    return datetime.now(tz=timezone.utc) - timedelta(days=days)


@router.get("/{cage_id}/state-budget")
def state_budget(
    cage_id: str,
    db: Annotated[Session, Depends(get_db)],
    _user: Annotated[User, Depends(get_current_user)],
    days: int = Query(7, ge=1, le=90),
    mouse: Optional[str] = None,
) -> dict[str, Any]:
    """Report 1 — time budget: per-mouse share of time Active / Rest / Sleep."""
    _ensure_cage(db, cage_id)
    iv = _sample_interval_s()
    q = (db.query(MouseStateSample.mouse_id, MouseStateSample.state,
                  func.count(MouseStateSample.id))
         .filter(MouseStateSample.cage_id == cage_id, MouseStateSample.ts >= _since(days)))
    if mouse:
        q = q.filter(MouseStateSample.mouse_id == mouse)
    rows = q.group_by(MouseStateSample.mouse_id, MouseStateSample.state).all()

    per_mouse: dict[str, dict[str, int]] = {}
    cage_states: dict[str, int] = {s: 0 for s in _CAM_STATES}
    for mid, state, n in rows:
        per_mouse.setdefault(mid, {})[state] = per_mouse.get(mid, {}).get(state, 0) + n
        cage_states[state] = cage_states.get(state, 0) + n

    def pack(states: dict[str, int]) -> dict[str, Any]:
        total = sum(states.values()) or 1
        return {
            "total_seconds": round(sum(states.values()) * iv, 1),
            "states": {s: {"seconds": round(states.get(s, 0) * iv, 1),
                           "hours": round(states.get(s, 0) * iv / 3600, 3),
                           "pct": round(100 * states.get(s, 0) / total, 1)}
                       for s in _CAM_STATES},
        }

    mice = [{"mouse_id": mid, **pack(states)} for mid, states in sorted(per_mouse.items())]
    return {"cage_id": cage_id, "range_days": days, "sample_interval_s": iv,
            "mice": mice, "cage_totals": pack(cage_states)}


@router.get("/{cage_id}/state-timeline")
def state_timeline(
    cage_id: str,
    db: Annotated[Session, Depends(get_db)],
    _user: Annotated[User, Depends(get_current_user)],
    date: Optional[str] = None,
    bucket: int = Query(300, ge=30, le=3600),
    mouse: Optional[str] = None,
) -> dict[str, Any]:
    """Report 2 — 24 h ethogram: dominant state per time bucket, per mouse."""
    _ensure_cage(db, cage_id)
    if date:
        day = datetime.strptime(date, "%Y-%m-%d").replace(tzinfo=timezone.utc)
    else:
        day = datetime.now(tz=timezone.utc).replace(hour=0, minute=0, second=0, microsecond=0)
    end = day + timedelta(days=1)
    q = (db.query(MouseStateSample.mouse_id, MouseStateSample.ts, MouseStateSample.state)
         .filter(MouseStateSample.cage_id == cage_id,
                 MouseStateSample.ts >= day, MouseStateSample.ts < end))
    if mouse:
        q = q.filter(MouseStateSample.mouse_id == mouse)
    rows = q.order_by(MouseStateSample.ts.asc()).all()

    base = day.timestamp()
    buckets: dict[str, dict[int, Counter]] = defaultdict(lambda: defaultdict(Counter))
    for mid, ts, state in rows:
        buckets[mid][int((ts.timestamp() - base) // bucket)][state] += 1
    mice = []
    for mid, bks in sorted(buckets.items()):
        segs = [{"t": datetime.fromtimestamp(base + bi * bucket, tz=timezone.utc)
                 .isoformat().replace("+00:00", "Z"),
                 "state": bks[bi].most_common(1)[0][0]} for bi in sorted(bks)]
        mice.append({"mouse_id": mid, "segments": segs})
    return {"cage_id": cage_id, "date": day.date().isoformat(),
            "bucket_seconds": bucket, "mice": mice}


@router.get("/{cage_id}/sleep-bouts")
def sleep_bouts(
    cage_id: str,
    db: Annotated[Session, Depends(get_db)],
    _user: Annotated[User, Depends(get_current_user)],
    days: int = Query(7, ge=1, le=90),
    mouse: Optional[str] = None,
    max_gap: int = Query(2, ge=0, le=60),
) -> dict[str, Any]:
    """Report 3 — sleep architecture: bouts, durations, fragmentation per mouse."""
    _ensure_cage(db, cage_id)
    iv = _sample_interval_s()
    q = (db.query(MouseStateSample.mouse_id, MouseStateSample.ts, MouseStateSample.state)
         .filter(MouseStateSample.cage_id == cage_id, MouseStateSample.ts >= _since(days)))
    if mouse:
        q = q.filter(MouseStateSample.mouse_id == mouse)
    rows = q.order_by(MouseStateSample.mouse_id.asc(), MouseStateSample.ts.asc()).all()

    by_mouse: dict[str, list] = defaultdict(list)
    for mid, ts, state in rows:
        by_mouse[mid].append((ts, state))
    gap_limit = (max_gap + 1) * iv
    out = []
    for mid, seq in sorted(by_mouse.items()):
        bouts: list[list] = []
        cur = None
        for ts, state in seq:
            if state == "Possible Sleep":
                if cur is None:
                    cur = [ts, ts]
                elif (ts - cur[1]).total_seconds() <= gap_limit:
                    cur[1] = ts
                else:
                    bouts.append(cur); cur = [ts, ts]
            elif cur is not None:
                bouts.append(cur); cur = None
        if cur is not None:
            bouts.append(cur)
        durs = [((b[1] - b[0]).total_seconds() + iv) / 60.0 for b in bouts]
        total = sum(durs)
        out.append({
            "mouse_id": mid, "bout_count": len(bouts),
            "total_sleep_min": round(total, 1),
            "mean_bout_min": round(statistics.mean(durs), 2) if durs else 0,
            "median_bout_min": round(statistics.median(durs), 2) if durs else 0,
            "max_bout_min": round(max(durs), 2) if durs else 0,
            "fragmentation_index": round(len(bouts) / (total / 60), 2) if total > 0 else 0,
            "bouts": [{"start": b[0].isoformat().replace("+00:00", "Z"),
                       "duration_min": round(d, 2)} for b, d in zip(bouts, durs)][:200],
        })
    return {"cage_id": cage_id, "range_days": days, "mice": out}


@router.get("/{cage_id}/state-circadian")
def state_circadian(
    cage_id: str,
    db: Annotated[Session, Depends(get_db)],
    _user: Annotated[User, Depends(get_current_user)],
    days: int = Query(7, ge=1, le=90),
    mouse: Optional[str] = None,
) -> dict[str, Any]:
    """Report 5 — circadian: active/rest/sleep fraction by hour-of-day (UTC)."""
    _ensure_cage(db, cage_id)
    hour_expr = func.strftime("%H", MouseStateSample.ts)
    q = (db.query(hour_expr, MouseStateSample.state, func.count(MouseStateSample.id))
         .filter(MouseStateSample.cage_id == cage_id, MouseStateSample.ts >= _since(days)))
    if mouse:
        q = q.filter(MouseStateSample.mouse_id == mouse)
    rows = q.group_by(hour_expr, MouseStateSample.state).all()

    hours = {h: {"active": 0, "rest": 0, "sleep": 0, "n": 0} for h in range(24)}
    for h, state, n in rows:
        try:
            b = hours[int(h)]
        except (ValueError, TypeError, KeyError):
            continue
        if state == "Wake / Active":
            b["active"] += n
        elif state == "Quiet / Rest":
            b["rest"] += n
        elif state == "Possible Sleep":
            b["sleep"] += n
        b["n"] += n
    out = []
    for h in range(24):
        b = hours[h]; tot = b["n"] or 1
        out.append({"hour": h, "active_pct": round(100 * b["active"] / tot, 1),
                    "rest_pct": round(100 * b["rest"] / tot, 1),
                    "sleep_pct": round(100 * b["sleep"] / tot, 1), "n": b["n"]})
    return {"cage_id": cage_id, "range_days": days, "timezone": "UTC", "hours": out}


@router.get("/{cage_id}/social-summary")
def social_summary(
    cage_id: str,
    db: Annotated[Session, Depends(get_db)],
    _user: Annotated[User, Depends(get_current_user)],
    days: int = Query(7, ge=1, le=90),
) -> dict[str, Any]:
    """Report 6a — group social state composition over days."""
    _ensure_cage(db, cage_id)
    day_expr = func.strftime("%Y-%m-%d", CageSocialSample.ts)
    rows = (db.query(day_expr, CageSocialSample.social_state,
                     func.count(CageSocialSample.id),
                     func.avg(CageSocialSample.mean_interanimal_dist))
            .filter(CageSocialSample.cage_id == cage_id, CageSocialSample.ts >= _since(days))
            .group_by(day_expr, CageSocialSample.social_state).all())
    days_map: dict[str, dict] = defaultdict(lambda: {"counts": defaultdict(int), "dist": []})
    for d, state, n, avgdist in rows:
        days_map[d]["counts"][state] += n
        if avgdist is not None:
            days_map[d]["dist"].append(avgdist)
    out = []
    for d in sorted(days_map):
        c = days_map[d]["counts"]; tot = sum(c.values()) or 1
        dist = days_map[d]["dist"]
        out.append({"date": d, "social": {k: round(100 * v / tot, 1) for k, v in c.items()},
                    "mean_interanimal_dist": round(sum(dist) / len(dist), 1) if dist else 0})
    return {"cage_id": cage_id, "range_days": days, "days": out}


@router.get("/{cage_id}/behavior-anomalies")
def behavior_anomalies(
    cage_id: str,
    db: Annotated[Session, Depends(get_db)],
    _user: Annotated[User, Depends(get_current_user)],
    days: int = Query(7, ge=1, le=90),
    mouse: Optional[str] = None,
) -> dict[str, Any]:
    """Report 6b — repetitive/stereotypy behaviour minutes per mouse per day."""
    _ensure_cage(db, cage_id)
    iv = _sample_interval_s()
    day_expr = func.strftime("%Y-%m-%d", MouseStateSample.ts)
    q = (db.query(day_expr, MouseStateSample.mouse_id, func.count(MouseStateSample.id))
         .filter(MouseStateSample.cage_id == cage_id, MouseStateSample.ts >= _since(days),
                 MouseStateSample.behavior == "Repetitive"))
    if mouse:
        q = q.filter(MouseStateSample.mouse_id == mouse)
    rows = q.group_by(day_expr, MouseStateSample.mouse_id).all()

    days_map: dict[str, list] = defaultdict(list)
    total_min = 0.0
    for d, mid, n in rows:
        rep_min = round(n * iv / 60, 2)
        total_min += rep_min
        days_map[d].append({"mouse_id": mid, "repetitive_min": rep_min})
    out = [{"date": d, "mice": sorted(days_map[d], key=lambda x: x["mouse_id"])}
           for d in sorted(days_map)]
    return {"cage_id": cage_id, "range_days": days, "days": out,
            "totals": {"repetitive_min": round(total_min, 1)}}
